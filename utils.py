import openpyxl
import PyPDF2
from docx import Document
import io
from googleapiclient.discovery import build
from google.oauth2 import service_account
from extensions import db
from models import Task, PDFAnalysis, GeneratedEmail
from chat_request import send_openai_request

def process_excel(file, add_to_db=True):
    workbook = openpyxl.load_workbook(file)
    sheet = workbook.active
    return extract_tasks_from_sheet(sheet, add_to_db)

def process_google_sheet(sheet_id, add_to_db=True):
    SCOPES = ['https://www.googleapis.com/auth/spreadsheets.readonly']
    SERVICE_ACCOUNT_FILE = 'google_credentials.json'
    
    try:
        credentials = service_account.Credentials.from_service_account_file(
            SERVICE_ACCOUNT_FILE, scopes=SCOPES)
        service = build('sheets', 'v4', credentials=credentials)
        sheet = service.spreadsheets()
        result = sheet.values().get(spreadsheetId=sheet_id, range='A1:Z').execute()
        values = result.get('values', [])
        
        if not values:
            return []
            
        header = values[0]
        task_index = None
        email_index = None
        recipient_index = None
        
        for i, cell in enumerate(header):
            if cell == 'Task':
                task_index = i
            elif cell == 'E-mail':
                email_index = i
            elif cell == 'Recipient':
                recipient_index = i
                
        if None in (task_index, email_index, recipient_index):
            raise ValueError("Required columns 'Task', 'E-mail', and 'Recipient' not found in the Google Sheet.")
            
        tasks = []
        for row in values[1:]:
            if len(row) > max(task_index, email_index, recipient_index):
                task_description = row[task_index]
                email = row[email_index]
                recipient = row[recipient_index]
                if task_description and email and recipient:
                    tasks.append({'description': task_description, 'email': email, 'recipient': recipient})
                    
        if add_to_db:
            save_tasks_to_db(tasks)
            
        return tasks
    except Exception as e:
        raise ValueError(f"Error processing Google Sheet: {str(e)}")

def process_word_document(file, add_to_db=True):
    try:
        doc = Document(file)
        tasks = []
        current_task = {}
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text.startswith("Task:"):
                if current_task:
                    if all(k in current_task for k in ['description', 'email', 'recipient']):
                        tasks.append(current_task)
                current_task = {'description': text[5:].strip()}
            elif text.startswith("Email:"):
                if current_task:
                    current_task['email'] = text[6:].strip()
            elif text.startswith("Recipient:"):
                if current_task:
                    current_task['recipient'] = text[10:].strip()
                    
        if current_task and all(k in current_task for k in ['description', 'email', 'recipient']):
            tasks.append(current_task)
            
        if not tasks:
            raise ValueError("No valid tasks found in the Word document. Format should be:\nTask: [description]\nEmail: [email]\nRecipient: [name]")
            
        if add_to_db:
            save_tasks_to_db(tasks)
            
        return tasks
    except Exception as e:
        raise ValueError(f"Error processing Word document: {str(e)}")

def extract_tasks_from_sheet(sheet, add_to_db=True):
    rows = list(sheet.iter_rows(min_row=1, values_only=True))
    
    if not rows:
        return []
    
    header = rows[0]
    task_index = None
    email_index = None
    recipient_index = None
    
    for i, cell in enumerate(header):
        if cell == 'Task':
            task_index = i
        elif cell == 'E-mail':
            email_index = i
        elif cell == 'Recipient':
            recipient_index = i
    
    if None in (task_index, email_index, recipient_index):
        raise ValueError("Required columns 'Task', 'E-mail', and 'Recipient' not found in the Excel file.")
    
    tasks = []
    for row in rows[1:]:
        if len(row) > max(task_index, email_index, recipient_index):
            task_description = row[task_index]
            email = row[email_index]
            recipient = row[recipient_index]
            if task_description and email and recipient:
                tasks.append({'description': task_description, 'email': email, 'recipient': recipient})
    
    if add_to_db:
        save_tasks_to_db(tasks)
    
    return tasks

def save_tasks_to_db(tasks):
    try:
        for task in tasks:
            db_task = Task(description=task['description'], email=task['email'], recipient=task['recipient'])
            db.session.add(db_task)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        print(f"Error saving tasks to database: {str(e)}")
        return []

def extract_document_text(file):
    """Extract text from either PDF or Word document."""
    if file.filename.endswith('.pdf'):
        reader = PyPDF2.PdfReader(file)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
    elif file.filename.endswith('.docx'):
        doc = Document(file)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    else:
        raise ValueError("Unsupported file format. Only PDF and DOCX files are supported.")

def analyze_documents(files):
    if len(files) != 2:
        raise ValueError("Exactly two documents are required for comparison.")
    
    texts = []
    for file in files:
        text = extract_document_text(file)
        texts.append(text)
    
    prompt = f"""Analyze and compare the following two documents for inconsistencies, logical fallacies, and unsupported statements. Highlight any discrepancies between them:

Document 1 ({files[0].filename}):
{texts[0][:4000]}

Document 2 ({files[1].filename}):
{texts[1][:4000]}

Provide a detailed analysis of the differences, inconsistencies, and potential issues found between these two documents."""

    analysis = send_openai_request(prompt)
    
    pdf_analysis = PDFAnalysis(filename=f"{files[0].filename}, {files[1].filename}", analysis=analysis)
    db.session.add(pdf_analysis)
    db.session.commit()
    
    return analysis

def generate_email(task, email, recipient):
    prompt = f"""Generate a professional email with the following details:
Task: {task}
Recipient: {recipient}
Recipient's Email: {email}

The email should be formal, clear, and concise. Include a brief introduction, details about the task, and a polite closing."""

    generated_email = send_openai_request(prompt)
    
    task_obj = Task.query.filter_by(description=task, email=email, recipient=recipient).first()
    if task_obj:
        email_obj = GeneratedEmail(task_id=task_obj.id, content=generated_email)
        db.session.add(email_obj)
        db.session.commit()
    else:
        print(f"Task not found for description: {task}, email: {email}, and recipient: {recipient}")
    
    return generated_email
